from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "D1D5DB") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False,
                   size: int = 11, color: str | None = None,
                   align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: int = 0, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # azul


def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E5E7EB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Gerador principal ─────────────────────────────────────────────────────────

def gerar_template_docx() -> bytes:
    """
    Gera um template .docx de Proposta Comercial com todas as variáveis
    do sistema ERP. O usuário baixa, personaliza no Word e faz re-upload.

    Variáveis disponíveis (docxtpl / Jinja2):
      Simples : {{ NUMERO }}, {{ TITULO }}, {{ CLIENTE }}, {{ VENDEDOR }},
                {{ DATA }}, {{ VALIDADE }}, {{ CUSTO_BASE }}, {{ SUBTOTAL }},
                {{ VALOR_VENDA }}, {{ OBSERVACOES }}, {{ EMAIL }},
                {{ TELEFONE }}, {{ ENDERECO }}, {{ STATUS }}, {{ EMPRESA }}
      Listas  : {{ PREMISSAS }} → campos NOME, TIPO, VALOR, CALCULADO
                {{ ITENS }}    → campos DESCRICAO, QUANTIDADE, VALOR
    """
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── CABEÇALHO ─────────────────────────────────────────────────────────────
    _add_paragraph(
        doc, "{{ EMPRESA }}",
        bold=True, size=20, color="1E3A8A",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_paragraph(
        doc, "PROPOSTA COMERCIAL",
        bold=True, size=14, color="374151",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_paragraph(
        doc, "Nº {{ NUMERO }}  ·  Status: {{ STATUS }}",
        size=10, color="6B7280",
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
    )
    _add_hr(doc)

    # ── DADOS DO ORÇAMENTO ────────────────────────────────────────────────────
    _add_section_title(doc, "Dados da Proposta")

    tbl_info = doc.add_table(rows=4, cols=4)
    tbl_info.style = "Table Grid"
    tbl_info.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Cm(3.5), Cm(8.5), Cm(3.5), Cm(4)]
    headers_vals = [
        ("Título:", "{{ TITULO }}", "Data:", "{{ DATA }}"),
        ("Cliente:", "{{ CLIENTE }}", "Validade:", "{{ VALIDADE }} dias"),
        ("Vendedor:", "{{ VENDEDOR }}", "E-mail:", "{{ EMAIL }}"),
        ("Telefone:", "{{ TELEFONE }}", "Endereço:", "{{ ENDERECO }}"),
    ]

    for row_i, (h1, v1, h2, v2) in enumerate(headers_vals):
        row = tbl_info.rows[row_i]
        for col_i, (text, is_header) in enumerate([(h1, True), (v1, False), (h2, True), (v2, False)]):
            cell = row.cells[col_i]
            cell.width = widths[col_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.bold = is_header
            _set_cell_borders(cell)
            if is_header:
                _set_cell_bg(cell, "F3F4F6")

    doc.add_paragraph()

    # ── PREMISSAS ─────────────────────────────────────────────────────────────
    _add_section_title(doc, "Premissas / Encargos")

    tbl_prem = doc.add_table(rows=3, cols=4)
    tbl_prem.style = "Table Grid"

    # Cabeçalho
    header_row = tbl_prem.rows[0]
    for col_i, txt in enumerate(["Premissa", "Tipo", "Valor (%/R$)", "Calculado"]):
        cell = header_row.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "1E40AF")
        _set_cell_borders(cell, "1E40AF")

    # Linha de loop — docxtpl usa {%tr for p in PREMISSAS %} na 1ª célula
    loop_row = tbl_prem.rows[1]
    loop_cells_prem = [
        ("{%tr for p in PREMISSAS %}{{ p.NOME }}", False),
        ("{{ p.TIPO }}", False),
        ("{{ p.VALOR }}", False),
        ("{{ p.CALCULADO }}", False),
    ]
    for col_i, (txt, _) in enumerate(loop_cells_prem):
        cell = loop_row.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        _set_cell_borders(cell)

    # Linha de fechamento do loop
    end_row = tbl_prem.rows[2]
    end_cells = [
        ("{%tr endfor %}", False),
        ("", False),
        ("", False),
        ("", False),
    ]
    for col_i, (txt, _) in enumerate(end_cells):
        cell = end_row.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        _set_cell_borders(cell)

    doc.add_paragraph()

    # ── ITENS / PRODUTOS ──────────────────────────────────────────────────────
    _add_section_title(doc, "Itens / Produtos")

    tbl_itens = doc.add_table(rows=3, cols=3)
    tbl_itens.style = "Table Grid"

    # Cabeçalho
    hdr_itens = tbl_itens.rows[0]
    for col_i, txt in enumerate(["Descrição", "Quantidade", "Valor"]):
        cell = hdr_itens.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "047857")
        _set_cell_borders(cell, "047857")

    # Linha de loop
    loop_itens = tbl_itens.rows[1]
    for col_i, txt in enumerate([
        "{%tr for i in ITENS %}{{ i.DESCRICAO }}",
        "{{ i.QUANTIDADE }}",
        "{{ i.VALOR }}",
    ]):
        cell = loop_itens.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        _set_cell_borders(cell)

    # Fechamento
    end_itens = tbl_itens.rows[2]
    for col_i, txt in enumerate(["{%tr endfor %}", "", ""]):
        cell = end_itens.cells[col_i]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        _set_cell_borders(cell)

    doc.add_paragraph()

    # ── TOTAIS ────────────────────────────────────────────────────────────────
    _add_section_title(doc, "Resumo Financeiro")

    tbl_totais = doc.add_table(rows=3, cols=2)
    tbl_totais.style = "Table Grid"

    totais_data = [
        ("Custo Base:", "{{ CUSTO_BASE }}"),
        ("Subtotal (com encargos):", "{{ SUBTOTAL }}"),
        ("VALOR DE VENDA:", "{{ VALOR_VENDA }}"),
    ]
    bgs = ["FFFFFF", "FFFFFF", "DBEAFE"]
    for row_i, (label, val) in enumerate(totais_data):
        row = tbl_totais.rows[row_i]
        is_total = row_i == 2

        label_cell = row.cells[0]
        val_cell = row.cells[1]

        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.size = Pt(10 if is_total else 9)
        lr.bold = is_total

        vp = val_cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vr = vp.add_run(val)
        vr.font.size = Pt(11 if is_total else 9)
        vr.bold = is_total
        if is_total:
            vr.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

        for cell in (label_cell, val_cell):
            _set_cell_bg(cell, bgs[row_i])
            _set_cell_borders(cell)

    doc.add_paragraph()
    _add_hr(doc)

    # ── OBSERVAÇÕES ───────────────────────────────────────────────────────────
    _add_section_title(doc, "Observações")
    _add_paragraph(doc, "{{ OBSERVACOES }}", size=9, color="374151", space_after=12)

    # ── ASSINATURA ────────────────────────────────────────────────────────────
    _add_section_title(doc, "Aceite")

    tbl_assin = doc.add_table(rows=2, cols=2)
    tbl_assin.style = "Table Grid"

    assin_labels = [
        ("Assinatura do Cliente:", "Assinatura do Responsável:"),
        ("Data: ___/___/______", "Vendedor: {{ VENDEDOR }}"),
    ]
    for row_i, (l1, l2) in enumerate(assin_labels):
        for col_i, txt in enumerate([l1, l2]):
            cell = tbl_assin.rows[row_i].cells[col_i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(16 if row_i == 0 else 4)
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.bold = row_i == 0
            _set_cell_borders(cell)
            if row_i == 0:
                _set_cell_bg(cell, "F9FAFB")

    # ── RODAPÉ ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "Proposta válida por {{ VALIDADE }} dias a partir de {{ DATA }}. "
        "Em caso de dúvidas, entre em contato: {{ EMAIL }} | {{ TELEFONE }}",
        size=8, color="9CA3AF",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=8,
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
